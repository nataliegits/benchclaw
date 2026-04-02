"""BenchClaw v2 — new feature renders and helpers.

Features:
  1. Export (.docx / .pdf)       — _export_docx, _export_pdf, render_save_export
  2. Protocol Diff + Audit       — render_diff_auditor
  3. LabClaw / DB Search         — render_labclaw (UniProt, ChEMBL, PubChem via REST)
  4. Reagent Cost Estimator      — render_reagent_cost
  5. My Protocols                — render_my_protocols
  6. OpenTrons Export            — render_opentrons
  7. Bench Vision                — render_bench_vision
"""

from __future__ import annotations

import base64
import difflib
import io
import json
import os

import requests
import streamlit as st

from benchclaw_db import db_save_protocol, db_user_protocols, db_delete_protocol

# ---------------------------------------------------------------------------
# Shared Claude client
# ---------------------------------------------------------------------------

@st.cache_resource
def get_client():
    import anthropic
    return anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])


# ---------------------------------------------------------------------------
# Feature 1 — Export helpers
# ---------------------------------------------------------------------------

def _export_docx(title: str, text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.add_heading(title, level=0)
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        # Detect markdown headings (## ...) and render as Word headings
        if block.startswith("### "):
            doc.add_heading(block[4:], level=3)
        elif block.startswith("## "):
            doc.add_heading(block[3:], level=2)
        elif block.startswith("# "):
            doc.add_heading(block[2:], level=1)
        else:
            p = doc.add_paragraph()
            # Bold **text** — simple inline pass
            remaining = block
            while "**" in remaining:
                pre, rest = remaining.split("**", 1)
                if "**" in rest:
                    bold_part, remaining = rest.split("**", 1)
                    if pre:
                        p.add_run(pre)
                    p.add_run(bold_part).bold = True
                else:
                    break
            if remaining:
                p.add_run(remaining)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _export_pdf(title: str, text: str) -> bytes:
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    # Title
    pdf.set_font("Helvetica", "B", 14)
    clean_title = title.encode("latin-1", errors="replace").decode("latin-1")
    pdf.multi_cell(0, 8, clean_title)
    pdf.ln(4)
    # Body — strip markdown syntax, clean for latin-1
    pdf.set_font("Helvetica", size=10)
    clean_text = (
        text
        .replace("**", "")
        .replace("__", "")
        .replace("# ", "")
        .replace("## ", "")
        .replace("### ", "")
        .encode("latin-1", errors="replace")
        .decode("latin-1")
    )
    pdf.multi_cell(0, 5, clean_text)
    return bytes(pdf.output())


def render_save_export(text: str, ptype: str, key_prefix: str) -> None:
    """Show export (.docx/.pdf) and save-to-share buttons after a generated result."""
    if not text or not text.strip():
        return

    st.divider()
    col_exp, col_save = st.columns([1, 1])

    # --- Export ---
    with col_exp:
        with st.expander("Export", expanded=False):
            title_input = st.text_input(
                "Document title",
                value=f"BenchClaw {ptype.title()}",
                key=f"{key_prefix}_export_title",
            )
            c1, c2 = st.columns(2)
            with c1:
                try:
                    docx_bytes = _export_docx(title_input, text)
                    st.download_button(
                        "Download .docx",
                        data=docx_bytes,
                        file_name=f"{title_input.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"{key_prefix}_dl_docx",
                    )
                except Exception as e:
                    st.error(f"Word export failed: {e}")
            with c2:
                try:
                    pdf_bytes = _export_pdf(title_input, text)
                    st.download_button(
                        "Download .pdf",
                        data=pdf_bytes,
                        file_name=f"{title_input.replace(' ', '_')}.pdf",
                        mime="application/pdf",
                        key=f"{key_prefix}_dl_pdf",
                    )
                except Exception as e:
                    st.error(f"PDF export failed: {e}")

    # --- Save & Share ---
    with col_save:
        with st.expander("Save & Share", expanded=False):
            user_id = st.session_state.get("user_id")
            if not user_id:
                st.info("Log in to save and share protocols.")
                return
            share_title = st.text_input(
                "Save as",
                value=f"My {ptype.title()}",
                key=f"{key_prefix}_share_title",
            )
            if st.button("Save & get share link", key=f"{key_prefix}_save_btn"):
                token = db_save_protocol(user_id, share_title, ptype, text)
                share_url = f"http://localhost:8501/?token={token}"
                st.success("Saved!")
                st.code(share_url, language=None)
                st.caption("Anyone with this link can view the protocol.")


# ---------------------------------------------------------------------------
# Feature 2 — Protocol Diff + Audit
# ---------------------------------------------------------------------------

_DIFF_AUDIT_SYSTEM = (
    "You are a senior molecular biologist reviewing a protocol revision. "
    "You are given a unified diff showing what changed between two versions of a lab protocol. "
    "Lines starting with + are additions, lines starting with - are removals. "
    "Audit the changes: are they improvements or regressions? "
    "Flag any removed safety steps, new parameter issues, or improvements. "
    "Be concise and specific."
)


def render_diff_auditor() -> None:
    st.header("Protocol Diff & Audit")
    st.caption("Compare two protocol versions — see what changed and get an AI audit of the delta.")

    col_a, col_b = st.columns(2)
    with col_a:
        v1 = st.text_area(
            "Version A (original)",
            height=280,
            placeholder="Paste original protocol…",
            key="diff_v1",
        )
    with col_b:
        v2 = st.text_area(
            "Version B (revised)",
            height=280,
            placeholder="Paste revised protocol…",
            key="diff_v2",
        )

    if st.button("Diff & Audit", type="primary", key="diff_btn"):
        if not v1.strip() or not v2.strip():
            st.warning("Please paste both protocol versions.")
            return

        old_lines = v1.splitlines(keepends=True)
        new_lines = v2.splitlines(keepends=True)
        diff_lines = list(
            difflib.unified_diff(old_lines, new_lines, fromfile="Version A", tofile="Version B")
        )
        diff_text = "".join(diff_lines)

        if not diff_text.strip():
            st.info("The two versions are identical.")
            return

        st.divider()
        st.subheader("Diff")
        st.code(diff_text, language="diff")

        # Download the diff
        st.download_button(
            "Download diff",
            data=diff_text,
            file_name="protocol_diff.diff",
            mime="text/plain",
            key="diff_download",
        )

        # Send only the changed lines to Claude
        delta_lines = [
            ln.rstrip("\n")
            for ln in diff_lines
            if ln.startswith(("+", "-", "@@")) and not ln.startswith(("+++", "---"))
        ]
        delta_for_claude = "\n".join(delta_lines)

        st.divider()
        st.subheader("AI Audit of Changes")

        def _stream():
            client = get_client()
            with client.messages.stream(
                model="claude-opus-4-6",
                max_tokens=2048,
                system=_DIFF_AUDIT_SYSTEM,
                messages=[{
                    "role": "user",
                    "content": (
                        "Audit these protocol changes. Lines with + are new; lines with - were removed.\n\n"
                        f"{delta_for_claude}"
                    ),
                }],
            ) as stream:
                for text in stream.text_stream:
                    yield text

        audit_text = st.write_stream(_stream())
        render_save_export(
            f"DIFF:\n{diff_text}\n\nAUDIT:\n{audit_text}",
            ptype="diff_audit",
            key_prefix="diff",
        )


# ---------------------------------------------------------------------------
# Feature 3 — LabClaw Database Search (UniProt / ChEMBL / PubChem via REST)
# ---------------------------------------------------------------------------

def _search_uniprot(query: str, limit: int = 5) -> list[dict]:
    resp = requests.get(
        "https://rest.uniprot.org/uniprotkb/search",
        params={
            "query": query,
            "format": "json",
            "size": limit,
            "fields": "accession,id,gene_names,protein_name,organism_name,cc_function",
        },
        timeout=15,
    )
    resp.raise_for_status()
    results = []
    for entry in resp.json().get("results", []):
        acc = entry.get("primaryAccession", "")
        # Protein name
        pn = entry.get("proteinDescription", {})
        rec = pn.get("recommendedName") or (pn.get("submissionNames") or [{}])[0]
        prot_name = rec.get("fullName", {}).get("value", "") if rec else ""
        # Genes
        genes = [
            g["geneName"]["value"]
            for g in entry.get("genes", [])[:3]
            if g.get("geneName")
        ]
        # Organism
        org = entry.get("organism", {}).get("scientificName", "")
        # Function
        func = ""
        for comment in entry.get("comments", []):
            if comment.get("commentType") == "FUNCTION":
                texts = comment.get("texts", [])
                if texts:
                    func = texts[0].get("value", "")[:300]
                    break
        results.append({
            "Accession": acc,
            "Protein": prot_name,
            "Genes": ", ".join(genes),
            "Organism": org,
            "Function": func,
            "URL": f"https://www.uniprot.org/uniprot/{acc}",
        })
    return results


def _search_chembl(query: str, limit: int = 5) -> list[dict]:
    resp = requests.get(
        "https://www.ebi.ac.uk/chembl/api/data/molecule.json",
        params={"pref_name__icontains": query, "limit": limit},
        timeout=15,
    )
    resp.raise_for_status()
    results = []
    for mol in resp.json().get("molecules", []):
        props = mol.get("molecule_properties") or {}
        cid = mol.get("molecule_chembl_id", "")
        results.append({
            "ChEMBL ID": cid,
            "Name": mol.get("pref_name", ""),
            "Type": mol.get("molecule_type", ""),
            "MW": props.get("mw_freebase", ""),
            "Formula": props.get("full_molformula", ""),
            "URL": f"https://www.ebi.ac.uk/chembl/compound_report_card/{cid}",
        })
    return results


def _search_pubchem(query: str) -> list[dict]:
    from urllib.parse import quote
    url = f"https://pubchem.ncbi.nlm.nih.gov/rest/pug/compound/name/{quote(query)}/JSON"
    resp = requests.get(url, timeout=15)
    if resp.status_code == 404:
        return []
    resp.raise_for_status()
    results = []
    for compound in resp.json().get("PC_Compounds", [])[:5]:
        cid = compound.get("id", {}).get("id", {}).get("cid", "")
        props = {}
        for prop in compound.get("props", []):
            label = prop.get("urn", {}).get("label", "")
            name_part = prop.get("urn", {}).get("name", "")
            val = prop.get("value", {})
            key = f"{label} {name_part}".strip()
            props[key] = val.get("sval") or val.get("fval") or val.get("ival") or ""
        results.append({
            "CID": str(cid),
            "IUPAC Name": props.get("IUPAC Name Preferred", props.get("IUPAC Name", "")),
            "Formula": props.get("Molecular Formula", ""),
            "MW": str(props.get("Molecular Weight", "")),
            "URL": f"https://pubchem.ncbi.nlm.nih.gov/compound/{cid}",
        })
    return results


def render_labclaw() -> None:
    st.header("Database Search")
    st.caption(
        "Search UniProt, ChEMBL, and PubChem — the same databases LabClaw wraps — directly from BenchClaw."
    )

    query = st.text_input(
        "Search term",
        placeholder="e.g. DNMT3A, 5-azacytidine, cyclosporin…",
        key="labclaw_query",
    )
    limit = st.slider("Results per database", min_value=1, max_value=10, value=5, key="labclaw_limit")

    if st.button("Search", type="primary", key="labclaw_btn"):
        if not query.strip():
            st.warning("Enter a search term.")
            return

        tabs = st.tabs(["UniProt", "ChEMBL", "PubChem"])

        with tabs[0]:
            st.subheader("UniProt — Protein Database")
            with st.spinner("Searching UniProt…"):
                try:
                    rows = _search_uniprot(query, limit)
                    if not rows:
                        st.info("No results.")
                    else:
                        for r in rows:
                            with st.expander(f"[{r['Accession']}] {r['Protein'] or r['Accession']}"):
                                st.markdown(f"**Genes:** {r['Genes']}")
                                st.markdown(f"**Organism:** {r['Organism']}")
                                if r["Function"]:
                                    st.markdown(f"**Function:** {r['Function']}")
                                st.markdown(f"[View on UniProt]({r['URL']})")
                except requests.RequestException as e:
                    st.warning(f"UniProt unavailable: {e}")

        with tabs[1]:
            st.subheader("ChEMBL — Bioactive Molecules")
            with st.spinner("Searching ChEMBL…"):
                try:
                    rows = _search_chembl(query, limit)
                    if not rows:
                        st.info("No results.")
                    else:
                        for r in rows:
                            with st.expander(f"[{r['ChEMBL ID']}] {r['Name'] or r['ChEMBL ID']}"):
                                st.markdown(f"**Type:** {r['Type']}")
                                st.markdown(f"**Formula:** {r['Formula']}  |  **MW:** {r['MW']}")
                                st.markdown(f"[View on ChEMBL]({r['URL']})")
                except requests.RequestException as e:
                    st.warning(f"ChEMBL unavailable: {e}")

        with tabs[2]:
            st.subheader("PubChem — Chemical Compounds")
            with st.spinner("Searching PubChem…"):
                try:
                    rows = _search_pubchem(query)
                    if not rows:
                        st.info("No results.")
                    else:
                        for r in rows:
                            with st.expander(f"[CID {r['CID']}] {r['IUPAC Name'][:60] or r['CID']}"):
                                st.markdown(f"**Formula:** {r['Formula']}  |  **MW:** {r['MW']}")
                                st.markdown(f"[View on PubChem]({r['URL']})")
                except requests.RequestException as e:
                    st.warning(f"PubChem unavailable: {e}")


# ---------------------------------------------------------------------------
# Feature 4 — Reagent Cost Estimator
# ---------------------------------------------------------------------------

VENDOR_HINTS: dict[str, dict] = {
    "antibody": {
        "vendors": ["Abcam", "Cell Signaling Technology", "Santa Cruz Biotech"],
        "range": "$200–600 per 100 μg",
    },
    "kit": {
        "vendors": ["Qiagen", "Thermo Fisher Scientific", "NEB"],
        "range": "$150–500 per kit",
    },
    "enzyme": {
        "vendors": ["NEB", "Thermo Fisher Scientific", "Promega"],
        "range": "$50–300 per 500 U",
    },
    "primer": {
        "vendors": ["IDT", "Sigma-Aldrich", "Eurofins Genomics"],
        "range": "$8–25 per oligo",
    },
    "bead": {
        "vendors": ["Thermo Fisher (Dynabeads)", "Miltenyi Biotec", "Cytiva"],
        "range": "$100–400 per mL",
    },
    "buffer": {
        "vendors": ["Sigma-Aldrich", "Thermo Fisher Scientific"],
        "range": "$20–80 per 500 mL",
    },
    "media": {
        "vendors": ["Gibco", "ATCC", "Sigma-Aldrich"],
        "range": "$20–60 per 500 mL",
    },
    "plasmid": {
        "vendors": ["Addgene", "IDT", "GenScript"],
        "range": "$65–200 per construct",
    },
    "chemical": {
        "vendors": ["Sigma-Aldrich", "Thermo Fisher Scientific", "Acros Organics"],
        "range": "$30–150 per gram",
    },
}

_REAGENT_SYSTEM = (
    "You are a lab supply expert. Extract all reagents, chemicals, kits, antibodies, "
    "and consumables mentioned in the lab protocol. "
    "Return ONLY a JSON array. Each item: "
    '{"name": "...", "quantity": "...", "category": "antibody|kit|enzyme|primer|bead|buffer|media|plasmid|chemical|other"}. '
    "No explanation. No markdown fences."
)


@st.cache_data(ttl=3600)
def _extract_reagents(protocol_text: str) -> list[dict]:
    client = get_client()
    msg = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=512,
        system=_REAGENT_SYSTEM,
        messages=[{"role": "user", "content": protocol_text[:4000]}],
    )
    raw = msg.content[0].text.strip()
    if raw.startswith("```"):
        raw = "\n".join(raw.split("\n")[1:-1])
    try:
        parsed = json.loads(raw)
        if isinstance(parsed, list):
            return parsed
        if isinstance(parsed, dict):
            return [parsed]
    except (json.JSONDecodeError, ValueError):
        pass
    return []


def render_reagent_cost() -> None:
    st.header("Reagent Cost Estimator")
    st.caption(
        "Paste a protocol — BenchClaw extracts all reagents and gives you "
        "typical price ranges and vendor suggestions."
    )

    protocol = st.text_area(
        "Protocol text",
        height=220,
        placeholder="Paste your protocol here…",
        key="reagent_protocol",
    )

    if st.button("Estimate Costs", type="primary", key="reagent_btn"):
        if not protocol.strip():
            st.warning("Please paste a protocol.")
            return

        with st.spinner("Extracting reagents…"):
            try:
                reagents = _extract_reagents(protocol)
            except Exception as e:
                st.error(f"Extraction failed: {e}")
                return

        if not reagents:
            st.warning("No reagents detected. Try a more detailed protocol.")
            return

        st.divider()
        st.subheader(f"Detected {len(reagents)} reagent(s)")

        rows = []
        for r in reagents:
            name = r.get("name", "Unknown")
            qty = r.get("quantity", "—")
            cat = r.get("category", "other").lower()
            hint = VENDOR_HINTS.get(cat, VENDOR_HINTS.get("chemical", {
                "vendors": ["Sigma-Aldrich", "Thermo Fisher Scientific"],
                "range": "Check vendor",
            }))
            rows.append({
                "Reagent": name,
                "Quantity": qty,
                "Category": cat.title(),
                "Est. Price": hint["range"],
                "Suggested Vendors": ", ".join(hint["vendors"]),
            })

        st.dataframe(rows, use_container_width=True, hide_index=True)

        # Rough total estimate
        st.caption(
            "Prices are typical list prices in USD. "
            "Academic discounts and bulk pricing can reduce costs significantly."
        )

        # Export
        try:
            import csv

            buf = io.StringIO()
            writer = csv.DictWriter(buf, fieldnames=list(rows[0].keys()))
            writer.writeheader()
            writer.writerows(rows)
            st.download_button(
                "Download as CSV",
                data=buf.getvalue(),
                file_name="reagent_costs.csv",
                mime="text/csv",
                key="reagent_csv",
            )
        except Exception:
            pass


# ---------------------------------------------------------------------------
# Feature 5 — My Protocols (saved protocol library)
# ---------------------------------------------------------------------------

def render_my_protocols() -> None:
    st.header("My Protocols")
    st.caption("Your saved protocols. Click a title to view; use share link to send to a colleague.")

    user_id = st.session_state.get("user_id")
    if not user_id:
        st.info("Log in to view saved protocols.")
        return

    protocols = db_user_protocols(user_id)

    if not protocols:
        st.info("No saved protocols yet. Use 'Save & Share' after generating or auditing a protocol.")
        return

    for row in protocols:
        col_title, col_link, col_del = st.columns([3, 2, 1])
        with col_title:
            with st.expander(f"**{row['title']}** — {row['ptype'].replace('_', ' ').title()}  ·  {row['created'][:10]}"):
                st.text(row["body"][:800] + ("…" if len(row["body"]) > 800 else ""))
        with col_link:
            share_url = f"http://localhost:8501/?token={row['token']}"
            st.code(share_url, language=None)
        with col_del:
            if st.button("Delete", key=f"del_{row['id']}"):
                db_delete_protocol(row["id"], user_id)
                st.rerun()


# ---------------------------------------------------------------------------
# Feature 6 — OpenTrons Protocol Export
# ---------------------------------------------------------------------------

_OPENTRONS_SYSTEM = """\
You are an expert Opentrons protocol engineer. Convert lab protocols into valid \
Python scripts for the Opentrons OT-2 using protocol_api version 2.16.

Rules:
- Always include a metadata dict with protocolName, author, description, apiLevel '2.16'
- Always define a run(protocol: protocol_api.ProtocolContext) function
- Use only real Opentrons labware names (e.g. opentrons_96_tiprack_300ul,
  corning_96_wellplate_360ul_flat, nest_12_reservoir_15ml,
  opentrons_24_tuberack_eppendorf_1.5ml_safelock_snapcap)
- Use only real pipette names (p20_single_gen2, p300_single_gen2,
  p1000_single_gen2, p20_multi_gen2, p300_multi_gen2)
- Add comments mapping each OT-2 step to the original protocol step number
- Where the protocol can't be fully automated, add a protocol.comment() call
  explaining what the scientist needs to do manually
- Output only the Python code, no explanation before or after
"""


def render_opentrons() -> None:
    st.header("OpenTrons Protocol Export")
    st.caption(
        "Paste any lab protocol and BenchClaw will convert it into a valid "
        "Python script for the Opentrons OT-2. Download and run it on the "
        "robot or paste it into the Opentrons simulator."
    )

    col_input, col_options = st.columns([3, 1])

    with col_input:
        protocol_text = st.text_area(
            "Protocol to convert",
            height=280,
            placeholder="Paste your protocol here...",
            key="ot_protocol_input",
        )

    with col_options:
        pipette_size = st.selectbox(
            "Primary pipette",
            ["p300_single_gen2", "p20_single_gen2", "p1000_single_gen2",
             "p300_multi_gen2", "p20_multi_gen2"],
            key="ot_pipette",
        )
        mount = st.selectbox("Mount", ["right", "left"], key="ot_mount")
        robot = st.selectbox("Robot", ["OT-2", "Flex (coming soon)"], key="ot_robot")
        st.caption(
            "Tip: the more specific your protocol steps and volumes, "
            "the more complete the generated code will be."
        )

    if st.button("Generate OT-2 Script", type="primary", key="ot_btn"):
        if not protocol_text.strip():
            st.warning("Please paste a protocol.")
            return
        if robot != "OT-2":
            st.warning("Flex support coming soon. Generating OT-2 script.")

        st.divider()
        st.subheader("Generated OT-2 Script")

        def _stream():
            client = get_client()
            with client.messages.stream(
                model="claude-opus-4-6",
                max_tokens=4096,
                system=_OPENTRONS_SYSTEM,
                messages=[{
                    "role": "user",
                    "content": (
                        f"Primary pipette: {pipette_size}, mount: {mount}\n\n"
                        f"PROTOCOL:\n{protocol_text}"
                    ),
                }],
            ) as stream:
                for text in stream.text_stream:
                    yield text

        ot_code = st.write_stream(_stream())
        ot_code = str(ot_code)

        st.divider()
        col_dl, col_sim = st.columns(2)
        with col_dl:
            st.download_button(
                "Download .py",
                data=ot_code,
                file_name="benchclaw_protocol.py",
                mime="text/x-python",
                key="ot_download",
            )
        with col_sim:
            st.link_button(
                "Open Opentrons Simulator",
                "https://designer.opentrons.com",
            )

        st.caption(
            "To test: open the Opentrons App, go to Protocol, upload the .py file. "
            "Or paste into the Opentrons Protocol Designer simulator."
        )

        render_save_export(ot_code, ptype="opentrons", key_prefix="ot")


# ---------------------------------------------------------------------------
# Feature 7 — Bench Vision (image upload + Claude vision)
# ---------------------------------------------------------------------------

_VISION_SYSTEM = """\
You are an expert lab scientist analyzing images from a biology research bench. \
When shown a lab image, you:
1. Describe exactly what you see (equipment, samples, reagents, labels, results)
2. Identify the likely experiment or technique being performed
3. Flag any visible issues: contamination, incorrect technique, mislabeled tubes,
   unusual gel bands, unhealthy cells, incorrect color reactions, safety concerns
4. If a protocol context is provided, note which step this image likely corresponds to
   and whether what you see matches expectations

Be specific and scientific. If you can read labels or text in the image, do so.
"""

_VISION_CATEGORIES = [
    "General lab photo",
    "Gel electrophoresis",
    "Cell culture / microscopy",
    "Western blot",
    "PCR / qPCR setup",
    "Reagent / tube labeling",
    "Lab notebook page",
    "Equipment / instrument",
    "Plate reader / assay",
    "Other",
]


def render_bench_vision() -> None:
    st.header("Bench Vision")
    st.caption(
        "Upload a photo from the bench. Claude will describe what it sees, "
        "identify the technique, and flag any issues. Optionally provide a "
        "protocol for context."
    )

    col_upload, col_context = st.columns([1, 1])

    with col_upload:
        uploaded = st.file_uploader(
            "Upload bench photo",
            type=["jpg", "jpeg", "png", "gif", "webp"],
            key="vision_upload",
        )
        if uploaded:
            st.image(uploaded, use_container_width=True)

        image_type = st.selectbox(
            "What kind of image is this?",
            _VISION_CATEGORIES,
            key="vision_category",
        )

    with col_context:
        protocol_context = st.text_area(
            "Protocol context (optional)",
            height=200,
            placeholder=(
                "Paste the relevant protocol or describe what step you're on. "
                "This helps Claude interpret the image more accurately."
            ),
            key="vision_protocol_context",
        )
        question = st.text_input(
            "Specific question (optional)",
            placeholder="e.g. Do these bands look right? Is this contaminated?",
            key="vision_question",
        )

    if st.button("Analyze Image", type="primary", key="vision_btn"):
        if not uploaded:
            st.warning("Please upload an image.")
            return

        uploaded.seek(0)
        image_bytes = uploaded.read()
        b64 = base64.standard_b64encode(image_bytes).decode()

        ext = uploaded.name.rsplit(".", 1)[-1].lower()
        media_type_map = {
            "jpg": "image/jpeg", "jpeg": "image/jpeg",
            "png": "image/png", "gif": "image/gif", "webp": "image/webp",
        }
        media_type = media_type_map.get(ext, "image/jpeg")

        user_content: list = [
            {
                "type": "image",
                "source": {"type": "base64", "media_type": media_type, "data": b64},
            },
        ]

        prompt_parts = [f"Image type: {image_type}"]
        if protocol_context.strip():
            prompt_parts.append(f"Protocol context:\n{protocol_context.strip()}")
        if question.strip():
            prompt_parts.append(f"Specific question: {question.strip()}")
        prompt_parts.append("Please analyze this lab image.")

        user_content.append({"type": "text", "text": "\n\n".join(prompt_parts)})

        st.divider()
        st.subheader("Analysis")

        def _stream():
            client = get_client()
            with client.messages.stream(
                model="claude-opus-4-6",
                max_tokens=2048,
                system=_VISION_SYSTEM,
                messages=[{"role": "user", "content": user_content}],
            ) as stream:
                for text in stream.text_stream:
                    yield text

        analysis = st.write_stream(_stream())
        render_save_export(str(analysis), ptype="vision_analysis", key_prefix="vision")
